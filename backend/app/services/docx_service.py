import io
import logging
from docx import Document

logger = logging.getLogger("uvicorn")

def extract_text_from_docx(docx_bytes: bytes) -> tuple[str, float]:
    """
    Extracts text from Word DOCX bytes using python-docx.
    Returns tuple of (extracted_text, confidence).
    """
    try:
        doc = Document(io.BytesIO(docx_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    paragraphs.append(" | ".join(row_text))

        full_text = "\n".join(paragraphs)
        confidence = 0.98 if len(full_text.strip()) > 30 else 0.5
        return full_text, confidence
    except Exception as e:
        logger.error(f"DOCX extraction error: {e}")
        return "", 0.0
